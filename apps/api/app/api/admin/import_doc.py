"""Document import — PDF / DOCX → vocabulary questions."""
import io
import re
import uuid
from typing import Annotated

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.deps import CurrentUser, require_role
from app.models.module import CefrLevel, Module, ModuleStatus, SourceKind, TopicType
from app.models.question import Question, QuestionKind
from app.models.user import UserRole

router = APIRouter(tags=["admin-import"])


class FinalizeBody(BaseModel):
    title: str
    cefr_level: str = "B1"
    questions: list[dict]


def _extract_words_pdf(data: bytes) -> list[dict]:
    """Extract vocabulary words from a PDF using pdfplumber."""
    try:
        import pdfplumber
        words = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    line = line.strip()
                    if not line or len(line) < 3:
                        continue
                    # Match: word (tab/multiple spaces) definition
                    parts = re.split(r"\t{1,}|\s{3,}", line, maxsplit=1)
                    if len(parts) == 2:
                        word, defn = parts[0].strip(), parts[1].strip()
                        if 1 <= len(word.split()) <= 3 and len(defn) >= 5:
                            words.append({"word": word, "definition": defn})
                    elif len(line.split()) <= 4:
                        # Standalone short phrase — treat as a word to define
                        words.append({"word": line, "definition": ""})
        return words
    except Exception as e:
        raise HTTPException(400, f"PDF extraction failed: {e}")


def _extract_words_docx(data: bytes) -> list[dict]:
    """Extract vocabulary words from a DOCX file."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        words = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    words.append({"word": cells[0], "definition": cells[1] if len(cells) > 1 else ""})
        if not words:
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                parts = re.split(r"\t{1,}|\s{3,}|:\s+", text, maxsplit=1)
                if len(parts) == 2 and len(parts[0].split()) <= 3:
                    words.append({"word": parts[0].strip(), "definition": parts[1].strip()})
        return words
    except Exception as e:
        raise HTTPException(400, f"DOCX extraction failed: {e}")


def _build_questions(words: list[dict]) -> list[dict]:
    """Build MCQ + fill + match questions from extracted vocabulary."""
    if not words:
        return []

    import random as _r
    questions = []
    words_with_def = [w for w in words if w.get("definition")]
    words_all = words

    chunk_size = 6

    if words_with_def:
        # Words have definitions — generate full MCQ + match questions
        for i in range(0, len(words_with_def), chunk_size):
            chunk = words_with_def[i: i + chunk_size]
            for entry in chunk[:3]:
                distractors = [w["definition"] for w in words_with_def if w["word"] != entry["word"]][:3]
                if len(distractors) < 3:
                    continue
                choices = [
                    {"id": "a", "label": entry["definition"]},
                    {"id": "b", "label": distractors[0]},
                    {"id": "c", "label": distractors[1]},
                    {"id": "d", "label": distractors[2]},
                ]
                _r.shuffle(choices)
                answer_id = next(c["id"] for c in choices if c["label"] == entry["definition"])
                questions.append({
                    "kind": "choice",
                    "prompt": f'What does "{entry["word"]}" mean?',
                    "payload": {"choices": choices, "answer": answer_id},
                    "explain": f'"{entry["word"]}" means: {entry["definition"]}',
                })

            if len(chunk) >= 4:
                questions.append({
                    "kind": "match",
                    "prompt": "Match each word to its correct definition.",
                    "payload": {
                        "pairs": [
                            {"left": e["word"], "right": e["definition"][:60]}
                            for e in chunk[:4]
                        ]
                    },
                })
    else:
        # Words without definitions — generate word recognition questions
        all_words = [w["word"] for w in words_all]
        for i in range(0, len(words_all), chunk_size):
            chunk = words_all[i: i + chunk_size]
            # Multiple choice: "which of these is a real English word from the document?"
            if len(chunk) >= 4:
                correct = chunk[0]["word"]
                fake = ["obfuscation", "synergize", "paradigmatic", "heuristic"]
                choices_labels = [correct] + fake[:3]
                _r.shuffle(choices_labels)
                choices = [{"id": chr(97 + j), "label": lbl} for j, lbl in enumerate(choices_labels)]
                answer_id = next(c["id"] for c in choices if c["label"] == correct)
                questions.append({
                    "kind": "choice",
                    "prompt": f'Which word appears in this document?',
                    "payload": {"choices": choices, "answer": answer_id},
                    "explain": f'"{correct}" is one of the key terms in this document.',
                })

            # Match question: group words (left) with themselves as the answer (placeholder)
            match_chunk = chunk[:4]
            if len(match_chunk) >= 4:
                shuffled = match_chunk[:]
                _r.shuffle(shuffled)
                questions.append({
                    "kind": "match",
                    "prompt": "Match each term from the document to its correct spelling.",
                    "payload": {
                        "pairs": [
                            {"left": e["word"].lower(), "right": e["word"]}
                            for e in match_chunk
                        ]
                    },
                })

    return questions


@router.post("/admin/import/document")
async def import_document(
    file: Annotated[UploadFile, File(...)],
    user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
    _: Annotated[None, Depends(require_role(UserRole.editor, UserRole.admin, UserRole.owner))],
):
    data = await file.read()
    filename = file.filename or ""

    if filename.lower().endswith(".pdf"):
        words = _extract_words_pdf(data)
    elif filename.lower().endswith((".doc", ".docx")):
        words = _extract_words_docx(data)
    else:
        raise HTTPException(400, "Only PDF and DOCX files are supported")

    if not words:
        raise HTTPException(422, "No vocabulary words could be extracted from the document")

    questions = _build_questions(words)

    # Archive original file to local storage for auditing
    try:
        from app.services.storage import put_object
        key = f"imports/{uuid.uuid4()}/{filename}"
        put_object(key, data, file.content_type or "application/octet-stream")
        source_blob = key
    except Exception:
        source_blob = None

    return {
        "title": filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title(),
        "word_count": len(words),
        "question_count": len(questions),
        "questions": questions,
        "source_blob": source_blob,
    }


@router.post("/admin/import/document/finalize", status_code=201)
async def finalize_import(
    body: FinalizeBody,
    user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
    _: Annotated[None, Depends(require_role(UserRole.editor, UserRole.admin, UserRole.owner))],
):
    try:
        cefr = CefrLevel(body.cefr_level)
    except ValueError:
        cefr = CefrLevel.B1

    module = Module(
        title=body.title,
        topic=TopicType.vocabulary,
        cefr_level=cefr,
        status=ModuleStatus.draft,
        created_by=user.id,
        source_kind=SourceKind.imported_pdf,
    )
    db.add(module)
    await db.flush()

    for i, qdata in enumerate(body.questions):
        q = Question(
            module_id=module.id,
            position=i,
            kind=QuestionKind(qdata["kind"]),
            prompt=qdata["prompt"],
            context=qdata.get("context"),
            sentence=qdata.get("sentence"),
            payload=qdata["payload"],
            explain=qdata.get("explain"),
        )
        db.add(q)

    await db.commit()
    return {"module_id": str(module.id), "title": module.title, "questions": len(body.questions)}
